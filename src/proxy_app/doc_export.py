"""Document export: turn AI-produced markdown into a downloadable .docx or .pdf,
optionally embedding a chart image.

Kept dependency-light and lazy: python-docx / reportlab / matplotlib are imported
only inside the function that needs them, so the proxy starts fast and a missing
optional lib degrades gracefully instead of breaking import.
"""
from __future__ import annotations

import io
import re
from typing import Any, Dict, List, Optional, Tuple


# ── Markdown parsing (a small, predictable subset) ──────────────────────────────
# Supported: # / ## / ### headings, - or * bullet lists, 1. numbered lists,
# | tables |, and plain paragraphs. Inline **bold** is flattened to plain text
# (kept simple on purpose — these documents are reports, not rich markup).

def _strip_inline(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text.strip()


def _parse_markdown(md: str, drop_title: str = "") -> List[Dict[str, Any]]:
    """Return a list of blocks: {'type': heading|para|bullet|ordered|table, ...}.

    If drop_title is given and the content starts with an H1 equal to it, that H1
    is skipped to avoid duplicating the document title.
    """
    blocks: List[Dict[str, Any]] = []
    md = (md or "").replace("\r\n", "\n")
    if drop_title:
        # Strip a leading "# <title>" that matches the document title.
        stripped = md.lstrip("\n")
        first = stripped.split("\n", 1)[0].strip()
        if first.lstrip("# ").strip().lower() == drop_title.strip().lower():
            md = stripped.split("\n", 1)[1] if "\n" in stripped else ""
    lines = md.split("\n")
    i = 0
    bullets: List[str] = []
    ordered: List[str] = []

    def flush_lists():
        nonlocal bullets, ordered
        if bullets:
            blocks.append({"type": "bullet", "items": bullets})
            bullets = []
        if ordered:
            blocks.append({"type": "ordered", "items": ordered})
            ordered = []

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            flush_lists()
            i += 1
            continue
        # table: a line with | and a following |---| separator
        if "|" in stripped and i + 1 < len(lines) and re.match(r"^\s*\|?[\s:|-]+\|?\s*$", lines[i + 1]):
            flush_lists()
            header = [c.strip() for c in stripped.strip("|").split("|")]
            rows = []
            i += 2
            while i < len(lines) and "|" in lines[i]:
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            blocks.append({"type": "table", "header": header, "rows": rows})
            continue
        m = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if m:
            flush_lists()
            blocks.append({"type": "heading", "level": len(m.group(1)), "text": _strip_inline(m.group(2))})
            i += 1
            continue
        m = re.match(r"^[-*]\s+(.*)$", stripped)
        if m:
            if ordered:
                flush_lists()
            bullets.append(_strip_inline(m.group(1)))
            i += 1
            continue
        m = re.match(r"^\d+[.)]\s+(.*)$", stripped)
        if m:
            if bullets:
                flush_lists()
            ordered.append(_strip_inline(m.group(1)))
            i += 1
            continue
        flush_lists()
        blocks.append({"type": "para", "text": _strip_inline(stripped)})
        i += 1
    flush_lists()
    return blocks


# ── Chart rendering (matplotlib, lazy, Agg backend) ─────────────────────────────

def _render_chart_png(chart: Dict[str, Any]) -> Optional[bytes]:
    """chart = {labels:[...], values:[...], title?, type?('bar'|'line'|'pie')}."""
    labels = chart.get("labels") or []
    values = chart.get("values") or []
    if not labels or not values or len(labels) != len(values):
        return None
    # Cap to prevent DoS: a malicious request with 100k labels would blow up
    # matplotlib memory + render time. 200 categories is more than any human
    # chart and beyond what the figure can render legibly.
    MAX_POINTS = 200
    if len(labels) > MAX_POINTS:
        labels = labels[:MAX_POINTS]
        values = values[:MAX_POINTS]
    try:
        import matplotlib
        matplotlib.use("Agg")  # no display, server-safe
        import matplotlib.pyplot as plt
    except Exception:
        return None
    try:
        values = [float(v) for v in values]
    except (TypeError, ValueError):
        return None
    kind = (chart.get("type") or "bar").lower()
    title = chart.get("title") or ""
    fig, ax = plt.subplots(figsize=(6.5, 3.6), dpi=130)
    colors = ["#6366f1", "#22c55e", "#3b82f6", "#f59e0b", "#ec4899", "#14b8a6"]
    try:
        if kind == "line":
            ax.plot(labels, values, color="#6366f1", marker="o", linewidth=2)
        elif kind == "pie":
            ax.pie(values, labels=labels, autopct="%1.0f%%",
                   colors=[colors[i % len(colors)] for i in range(len(values))])
        else:
            ax.bar(labels, values, color=[colors[i % len(colors)] for i in range(len(values))])
        if title:
            ax.set_title(title, fontsize=12, fontweight="bold")
        if kind != "pie":
            ax.spines["top"].set_visible(False)
            ax.spines["right"].set_visible(False)
            plt.xticks(rotation=30, ha="right", fontsize=8)
        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format="png", bbox_inches="tight")
        plt.close(fig)
        return buf.getvalue()
    except Exception:
        plt.close(fig)
        return None


# ── DOCX ────────────────────────────────────────────────────────────────────────

def build_docx(title: str, content_md: str, chart: Optional[Dict[str, Any]] = None) -> bytes:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor

    doc = Document()
    if title:
        h = doc.add_heading(title, level=0)
    for block in _parse_markdown(content_md, title):
        t = block["type"]
        if t == "heading":
            doc.add_heading(block["text"], level=min(block["level"], 4))
        elif t == "para":
            doc.add_paragraph(block["text"])
        elif t == "bullet":
            for it in block["items"]:
                doc.add_paragraph(it, style="List Bullet")
        elif t == "ordered":
            for it in block["items"]:
                doc.add_paragraph(it, style="List Number")
        elif t == "table":
            cols = max(len(block["header"]), *(len(r) for r in block["rows"])) if block["rows"] else len(block["header"])
            table = doc.add_table(rows=1, cols=cols)
            table.style = "Light Grid Accent 1"
            for idx, cell in enumerate(block["header"]):
                if idx < cols:
                    table.rows[0].cells[idx].text = cell
            for row in block["rows"]:
                cells = table.add_row().cells
                for idx, val in enumerate(row):
                    if idx < cols:
                        cells[idx].text = val
    if chart:
        png = _render_chart_png(chart)
        if png:
            doc.add_paragraph()
            doc.add_picture(io.BytesIO(png), width=Inches(6.0))
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# ── PDF (reportlab) ──────────────────────────────────────────────────────────────

def build_pdf(title: str, content_md: str, chart: Optional[Dict[str, Any]] = None) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem,
        Table, TableStyle, Image,
    )

    styles = getSampleStyleSheet()
    out = io.BytesIO()
    docp = SimpleDocTemplate(out, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm,
                             leftMargin=2 * cm, rightMargin=2 * cm)
    flow = []
    if title:
        flow.append(Paragraph(title, styles["Title"]))
        flow.append(Spacer(1, 12))
    for block in _parse_markdown(content_md, title):
        t = block["type"]
        if t == "heading":
            flow.append(Paragraph(block["text"], styles[f"Heading{min(block['level'],3)}"]))
        elif t == "para":
            flow.append(Paragraph(block["text"], styles["BodyText"]))
            flow.append(Spacer(1, 6))
        elif t in ("bullet", "ordered"):
            items = [ListItem(Paragraph(it, styles["BodyText"])) for it in block["items"]]
            flow.append(ListFlowable(items, bulletType="bullet" if t == "bullet" else "1"))
            flow.append(Spacer(1, 6))
        elif t == "table":
            data = [block["header"]] + block["rows"]
            tbl = Table(data, hAlign="LEFT")
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#6366f1")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f4f4f8")]),
                ("PADDING", (0, 0), (-1, -1), 5),
            ]))
            flow.append(tbl)
            flow.append(Spacer(1, 8))
    if chart:
        png = _render_chart_png(chart)
        if png:
            flow.append(Spacer(1, 10))
            flow.append(Image(io.BytesIO(png), width=15 * cm, height=8 * cm))
    docp.build(flow)
    return out.getvalue()
